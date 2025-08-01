import os
import logging
import asyncio
from typing import List, Optional, Dict, Any
from contextlib import asynccontextmanager

import uvicorn
from fastapi import FastAPI, HTTPException, UploadFile, File, Depends
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import Response
from pydantic import BaseModel, Field
from prometheus_client import Counter, Histogram, Gauge, generate_latest
import yaml

# LlamaIndex imports
from llama_index.core import VectorStoreIndex, Document, Settings
from llama_index.core.node_parser import SentenceSplitter
from llama_index.vector_stores.qdrant import QdrantVectorStore
from llama_index.embeddings.huggingface import HuggingFaceEmbedding
from llama_index.llms.ollama import Ollama

# Qdrant client
from qdrant_client import QdrantClient
from qdrant_client.models import Distance, VectorParams

# Database
import sqlalchemy as sa
from sqlalchemy.ext.declarative import declarative_base
from sqlalchemy.orm import sessionmaker, Session

# Document processing
import aiofiles
from pypdf import PdfReader
from docx import Document as DocxDocument
import pandas as pd
from bs4 import BeautifulSoup
import io

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Load configuration
with open("/app/config.yaml", "r") as f:
    config = yaml.safe_load(f)

# Database setup
DATABASE_URL = os.getenv('DATABASE_URL', 'postgresql://postgres:password@postgres:5432/aibox')

engine = sa.create_engine(DATABASE_URL)
SessionLocal = sessionmaker(autocommit=False, autoflush=False, bind=engine)
Base = declarative_base()

# Database models
class DocumentModel(Base):
    __tablename__ = "documents"

    id = sa.Column(sa.Integer, primary_key=True, index=True)
    filename = sa.Column(sa.String, index=True)
    content_hash = sa.Column(sa.String, unique=True, index=True)
    document_metadata = sa.Column(sa.JSON)
    created_at = sa.Column(sa.DateTime, default=sa.func.now())
    updated_at = sa.Column(sa.DateTime, default=sa.func.now(), onupdate=sa.func.now())

# Pydantic models
class DocumentUploadResponse(BaseModel):
    document_id: str
    filename: str
    status: str
    message: str

class QueryRequest(BaseModel):
    query: str = Field(..., description="The query to search for")
    top_k: int = Field(default=5, description="Number of top results to return")
    similarity_threshold: float = Field(default=0.7, description="Minimum similarity threshold")

class QueryResponse(BaseModel):
    answer: str
    sources: List[Dict[str, Any]]
    metadata: Dict[str, Any]

class HealthResponse(BaseModel):
    status: str
    services: Dict[str, str]

# Global variables
vector_store = None
index = None
query_engine = None
qdrant_client = None

@asynccontextmanager
async def lifespan(app: FastAPI):
    """Initialize services on startup"""
    global vector_store, index, query_engine, qdrant_client

    try:
        # Initialize database
        Base.metadata.create_all(bind=engine)
        logger.info("Database initialized")

        # Initialize Qdrant client
        qdrant_client = QdrantClient(
            host=os.getenv("QDRANT_HOST", "qdrant"),
            port=int(os.getenv("QDRANT_PORT", "6333"))
        )

        # Create collection if it doesn't exist
        collection_name = config["vector_store"]["collection_name"]
        try:
            qdrant_client.get_collection(collection_name)
            logger.info(f"Collection '{collection_name}' already exists")
        except Exception as e:
            if "doesn't exist" in str(e) or "Not found" in str(e):
                logger.info(f"Collection '{collection_name}' not found, creating...")
                qdrant_client.create_collection(
                    collection_name=collection_name,
                    vectors_config=VectorParams(
                        size=config["embeddings"]["dimension"],
                        distance=Distance.COSINE
                    )
                )
                logger.info(f"Created collection '{collection_name}'")
            else:
                logger.error(f"Error checking collection: {e}")
                raise

        # Initialize vector store
        vector_store = QdrantVectorStore(
            client=qdrant_client,
            collection_name=collection_name
        )

        # Initialize LLM
        llm = Ollama(
            model=config["llm"]["model"],
            base_url=os.getenv("OLLAMA_API_BASE", "http://ollama:11434"),
            temperature=config["llm"]["temperature"]
        )

        # Initialize embeddings with fallback
        offline_mode = os.getenv("TRANSFORMERS_OFFLINE", "0") == "1"

        if offline_mode:
            logger.info("Running in offline mode, using cached models only")

        try:
            embed_model = HuggingFaceEmbedding(
                model_name=config["embeddings"]["model"],
                cache_folder="/app/embeddings_cache"
            )
            logger.info(f"Successfully loaded embedding model: {config['embeddings']['model']}")
        except Exception as e:
            logger.warning(f"Failed to load primary embedding model: {e}")
            logger.info("Falling back to alternative embedding models")

            # Try fallback models in order of preference
            fallback_models = [
                "sentence-transformers/all-mpnet-base-v2",
                "BAAI/bge-small-en",
                "sentence-transformers/all-MiniLM-L6-v2"
            ]

            embed_model = None
            for fallback_model in fallback_models:
                try:
                    embed_model = HuggingFaceEmbedding(
                        model_name=fallback_model,
                        cache_folder="/app/embeddings_cache"
                    )
                    logger.info(f"Successfully loaded fallback embedding model: {fallback_model}")
                    break
                except Exception as fallback_error:
                    logger.warning(f"Failed to load {fallback_model}: {fallback_error}")
                    continue

            if embed_model is None:
                logger.error("All embedding models failed to load")
                raise Exception("No embedding model could be loaded")

        # Configure global settings
        Settings.llm = llm
        Settings.embed_model = embed_model
        Settings.chunk_size = config["text_processing"]["chunk_size"]
        Settings.chunk_overlap = config["text_processing"]["chunk_overlap"]

        # Initialize index
        index = VectorStoreIndex.from_vector_store(vector_store)
        query_engine = index.as_query_engine(
            similarity_top_k=config["retrieval"]["top_k"],
            response_mode="tree_summarize"
        )

        logger.info("RAG service initialized successfully")

    except Exception as e:
        logger.error(f"Failed to initialize RAG service: {e}")
        raise

    yield

    # Cleanup
    logger.info("Shutting down RAG service")

# Initialize FastAPI app
app = FastAPI(
    title="AI Box RAG Service",
    description="Retrieval-Augmented Generation service for AI Box",
    version="1.0.0",
    lifespan=lifespan
)

# Add CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Dependency to get database session
def get_db():
    db = SessionLocal()
    try:
        yield db
    finally:
        db.close()

# Document processing functions
async def process_pdf(file_content: bytes) -> str:
    """Extract text from PDF"""
    try:
        pdf_reader = PdfReader(io.BytesIO(file_content))
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text() + "\n"
        return text
    except Exception as e:
        logger.error(f"Error processing PDF: {e}")
        raise HTTPException(status_code=400, detail=f"Error processing PDF: {e}")

async def process_docx(file_content: bytes) -> str:
    """Extract text from DOCX"""
    try:
        doc = DocxDocument(io.BytesIO(file_content))
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text
    except Exception as e:
        logger.error(f"Error processing DOCX: {e}")
        raise HTTPException(status_code=400, detail=f"Error processing DOCX: {e}")

async def process_txt(file_content: bytes) -> str:
    """Extract text from TXT"""
    try:
        return file_content.decode('utf-8')
    except UnicodeDecodeError:
        return file_content.decode('latin-1')

async def process_html(file_content: bytes) -> str:
    """Extract text from HTML"""
    try:
        soup = BeautifulSoup(file_content, 'html.parser')
        return soup.get_text()
    except Exception as e:
        logger.error(f"Error processing HTML: {e}")
        raise HTTPException(status_code=400, detail=f"Error processing HTML: {e}")

# API endpoints
@app.get("/health", response_model=HealthResponse)
async def health_check():
    """Health check endpoint"""
    services = {}

    # Check Qdrant
    try:
        collections = qdrant_client.get_collections()
        services["qdrant"] = "healthy"
        logger.debug(f"Qdrant collections: {[c.name for c in collections.collections]}")
    except Exception as e:
        logger.error(f"Qdrant health check failed: {e}")
        services["qdrant"] = "unhealthy"

    # Check Ollama
    try:
        import requests
        response = requests.get(f"{os.getenv('OLLAMA_API_BASE', 'http://ollama:11434')}/api/tags", timeout=5)
        services["ollama"] = "healthy" if response.status_code == 200 else "unhealthy"
    except Exception:
        services["ollama"] = "unhealthy"

    # Check Database
    try:
        with engine.connect() as conn:
            conn.execute(sa.text("SELECT 1"))
        services["database"] = "healthy"
    except Exception:
        services["database"] = "unhealthy"

    status = "healthy" if all(s == "healthy" for s in services.values()) else "partial"

    return HealthResponse(status=status, services=services)

@app.post("/documents/upload", response_model=DocumentUploadResponse)
async def upload_document(
    file: UploadFile = File(...),
    db: Session = Depends(get_db)
):
    """Upload and process a document"""
    global index

    if not file.filename:
        raise HTTPException(status_code=400, detail="No file provided")

    # Read file content
    content = await file.read()

    # Process based on file type
    file_extension = file.filename.lower().split('.')[-1]

    try:
        if file_extension == 'pdf':
            text = await process_pdf(content)
        elif file_extension == 'docx':
            text = await process_docx(content)
        elif file_extension == 'txt':
            text = await process_txt(content)
        elif file_extension in ['html', 'htm']:
            text = await process_html(content)
        else:
            raise HTTPException(status_code=400, detail=f"Unsupported file type: {file_extension}")

        if not text.strip():
            raise HTTPException(status_code=400, detail="No text content found in document")

        # Create document hash
        import hashlib
        content_hash = hashlib.sha256(content).hexdigest()

        # Check if document already exists
        existing_doc = db.query(DocumentModel).filter(DocumentModel.content_hash == content_hash).first()
        if existing_doc:
            return DocumentUploadResponse(
                document_id=str(existing_doc.id),
                filename=file.filename,
                status="already_exists",
                message="Document already processed"
            )

        # Create document record
        doc_record = DocumentModel(
            filename=file.filename,
            content_hash=content_hash,
            document_metadata={
                "file_type": file_extension,
                "file_size": len(content),
                "content_length": len(text)
            }
        )
        db.add(doc_record)
        db.commit()
        db.refresh(doc_record)

        # Process text and add to vector store
        document = Document(
            text=text,
            metadata={
                "filename": file.filename,
                "document_id": str(doc_record.id),
                "file_type": file_extension
            }
        )

        # Add to index
        index.insert(document)

        logger.info(f"Successfully processed document: {file.filename}")

        return DocumentUploadResponse(
            document_id=str(doc_record.id),
            filename=file.filename,
            status="success",
            message="Document processed and indexed successfully"
        )

    except Exception as e:
        logger.error(f"Error processing document {file.filename}: {e}")
        db.rollback()
        raise HTTPException(status_code=500, detail=f"Error processing document: {e}")

@app.post("/query", response_model=QueryResponse)
async def query_documents(request: QueryRequest):
    """Query the document index"""
    global query_engine

    if not query_engine:
        raise HTTPException(status_code=503, detail="RAG service not initialized")

    try:
        # Execute query
        response = query_engine.query(request.query)

        # Extract sources
        sources = []
        if hasattr(response, 'source_nodes'):
            for node in response.source_nodes:
                sources.append({
                    "content": node.text[:500] + "..." if len(node.text) > 500 else node.text,
                    "metadata": node.metadata,
                    "score": getattr(node, 'score', None)
                })

        # Prepare metadata
        metadata = {
            "query": request.query,
            "top_k": request.top_k,
            "num_sources": len(sources)
        }

        return QueryResponse(
            answer=str(response),
            sources=sources,
            metadata=metadata
        )

    except Exception as e:
        logger.error(f"Error querying documents: {e}")
        raise HTTPException(status_code=500, detail=f"Error querying documents: {e}")

@app.get("/documents")
async def list_documents(db: Session = Depends(get_db)):
    """List all uploaded documents"""
    documents = db.query(DocumentModel).all()
    return [
        {
            "id": doc.id,
            "filename": doc.filename,
            "metadata": doc.document_metadata,
            "created_at": doc.created_at,
            "updated_at": doc.updated_at
        }
        for doc in documents
    ]

@app.delete("/documents/{document_id}")
async def delete_document(document_id: int, db: Session = Depends(get_db)):
    """Delete a document from the index"""
    document = db.query(DocumentModel).filter(DocumentModel.id == document_id).first()
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")

    try:
        # Note: In a production system, you would also need to remove
        # the corresponding vectors from the vector store
        db.delete(document)
        db.commit()

        return {"message": "Document deleted successfully"}

    except Exception as e:
        logger.error(f"Error deleting document: {e}")
        db.rollback()
        raise HTTPException(status_code=500, detail=f"Error deleting document: {e}")

@app.get("/metrics")
async def metrics():
    """Prometheus metrics endpoint"""
    from prometheus_client import generate_latest
    from fastapi.responses import Response
    return Response(generate_latest(), media_type="text/plain")

if __name__ == "__main__":
    uvicorn.run(
        "app:app",
        host="0.0.0.0",
        port=8001,
        reload=False,
        workers=4,
        log_level="info"
    )